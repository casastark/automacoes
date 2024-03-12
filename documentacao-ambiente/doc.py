import boto3
from docx import Document
from docx.shared import Pt

# Inicializa os clientes da AWS para os serviços EC2, S3, CloudFront, ELB, RDS, EKS, AWS Backup, NAT Gateways, Lambda usando o perfil padrão
session = boto3.Session()
ec2_client = session.client('ec2')
s3_client = session.client('s3')
cloudfront_client = session.client('cloudfront')
elbv2_client = session.client('elbv2')
rds_client = session.client('rds')
eks_client = session.client('eks')
backup_client = session.client('backup')
lambda_client = session.client('lambda')

# Inicializa o documento
doc = Document()

# Função para obter o nome de um recurso a partir de suas tags
def get_resource_name(resource):
    for tag in resource.get('Tags', []):
        if tag['Key'] == 'Name':
            return tag['Value']
    return 'N/A'

# Função para adicionar um título e uma lista de informações ao documento
def add_section(title, info_list, level):
    run = doc.add_heading(title, level).bold = True
    for info in info_list:
        doc.add_paragraph(f'• {info}')
    doc.add_paragraph()  # Adiciona um parágrafo em branco

# Obtém informações sobre VPCs e recursos relacionados
vpcs_info = {}
vpcs_response = ec2_client.describe_vpcs()
for vpc in vpcs_response['Vpcs']:
    vpc_id = vpc['VpcId']
    vpcs_info[vpc_id] = {
        'VPC ID': vpc_id,
        'VPC Name': get_resource_name(vpc),
        'CIDR Block': vpc['CidrBlock'],
        'Subnets': [],
        'Security Groups': [],
        'Instances': [],
        'Public IPs': [],
        'NAT Gateways': []
    }

    # Obtém informações sobre subnets relacionadas à VPC
    subnets_response = ec2_client.describe_subnets(Filters=[{'Name': 'vpc-id', 'Values': [vpc_id]}])
    for subnet in subnets_response['Subnets']:
        subnet_info = f'Subnet ID: {subnet["SubnetId"]} - Subnet Name: {get_resource_name(subnet)} - CIDR Block: {subnet["CidrBlock"]}'
        vpcs_info[vpc_id]['Subnets'].append(subnet_info)

    # Obtém informações sobre security groups relacionados à VPC
    security_groups_response = ec2_client.describe_security_groups(Filters=[{'Name': 'vpc-id', 'Values': [vpc_id]}])
    for sg in security_groups_response['SecurityGroups']:
        sg_info = f'Security Group ID: {sg["GroupId"]} - Security Group Name: {get_resource_name(sg)} - Description: {sg["Description"]}'
        vpcs_info[vpc_id]['Security Groups'].append(sg_info)

    # Obtém informações sobre instâncias EC2 relacionadas à VPC
    instances_response = ec2_client.describe_instances(Filters=[{'Name': 'vpc-id', 'Values': [vpc_id]}])
    for reservation in instances_response['Reservations']:
        for instance in reservation['Instances']:
            instance_info = f'Instance ID: {instance["InstanceId"]} - Instance Name: {get_resource_name(instance)} - Instance Type: {instance["InstanceType"]} - Private IP: {instance.get("PrivateIpAddress", "N/A")} - Public IP: {instance.get("PublicIpAddress", "N/A")}'
            vpcs_info[vpc_id]['Instances'].append(instance_info)
            if instance.get('PublicIpAddress'):
                vpcs_info[vpc_id]['Public IPs'].append(instance.get('PublicIpAddress'))

    # Obtém informações sobre NAT Gateways relacionados à VPC
    nat_gateways_response = ec2_client.describe_nat_gateways(Filters=[{'Name': 'vpc-id', 'Values': [vpc_id]}])
    for nat_gateway in nat_gateways_response['NatGateways']:
        nat_gateway_info = f'NAT Gateway ID: {nat_gateway["NatGatewayId"]} - Subnet ID: {nat_gateway["SubnetId"]} - Public IP: {nat_gateway.get("NatGatewayAddresses", [{}])[0].get("PublicIp", "N/A")}'
        vpcs_info[vpc_id]['NAT Gateways'].append(nat_gateway_info)

# Obtém informações sobre CloudFront distributions
cloudfront_info = []
cloudfront_response = cloudfront_client.list_distributions()
for distribution in cloudfront_response['DistributionList'].get('Items', []):
    cloudfront_info.append(f'Distribution ID: {distribution["Id"]} - Domain Name: {distribution["DomainName"]}')

# Obtém informações sobre Load Balancers (ELB)
elb_info = []
elb_response = elbv2_client.describe_load_balancers()
for elb in elb_response['LoadBalancers']:
    elb_info.append(f'Load Balancer ARN: {elb["LoadBalancerArn"]} - DNS Name: {elb["DNSName"]}')

# Obtém informações sobre instâncias RDS
rds_info = []
rds_response = rds_client.describe_db_instances()
for db_instance in rds_response['DBInstances']:
    rds_info.append(f'DB Instance Identifier: {db_instance["DBInstanceIdentifier"]} - Engine: {db_instance["Engine"]} - Endpoint: {db_instance["Endpoint"]["Address"]}')

# Obtém informações sobre clusters EKS
eks_info = []
eks_response = eks_client.list_clusters()
for cluster_name in eks_response['clusters']:
    cluster_info = eks_client.describe_cluster(name=cluster_name)
    eks_info.append(f'Cluster Name: {cluster_info["cluster"]["name"]} - Status: {cluster_info["cluster"]["status"]} - Endpoint: {cluster_info["cluster"]["endpoint"]}')

# Obtém informações sobre backups do AWS Backup
backup_info = []
backups_response = backup_client.list_backup_plans()
for backup in backups_response['BackupPlansList']:
    backup_id = backup.get('BackupPlanId', 'N/A')
    backup_name = backup.get('BackupPlanName', 'N/A')
    backup_info.append(f'Backup Plan ID: {backup_id} - Backup Plan Name: {backup_name}')

# Obtém informações sobre Lambda
lambda_info = []
lambda_response = lambda_client.list_functions()
for function in lambda_response['Functions']:
    lambda_info.append(f'Function Name: {function["FunctionName"]} - Runtime: {function.get("Runtime", "N/A")} - Memory Size: {function.get("MemorySize", "N/A")} MB')

# Obtém informações sobre S3 (apenas nome do bucket e número de objetos)
s3_info = []
s3_response = s3_client.list_buckets()
for bucket in s3_response['Buckets']:
    bucket_name = bucket['Name']
    objects_count_response = s3_client.list_objects_v2(Bucket=bucket_name)
    objects_count = len(objects_count_response.get('Contents', []))
    s3_info.append(f'Bucket Name: {bucket_name} - Number of Objects: {objects_count}')

# Adiciona as informações ao documento
for vpc_id, vpc_info in vpcs_info.items():
    add_section(f'VPC ID: {vpc_info["VPC ID"]} - VPC Name: {vpc_info["VPC Name"]} - CIDR Block: {vpc_info["CIDR Block"]}', vpc_info['Subnets'], 2)
    add_section('Security Groups:', vpc_info['Security Groups'], 3)
    add_section('Instances:', vpc_info['Instances'], 3)
    add_section('Public IPs:', vpc_info['Public IPs'], 3)
    add_section('NAT Gateways:', vpc_info['NAT Gateways'], 3)

add_section('CloudFront Distributions:', cloudfront_info, 2)
add_section('Load Balancers (ELB):', elb_info, 2)
add_section('Instâncias RDS:', rds_info, 2)
add_section('Clusters EKS:', eks_info, 2)
add_section('Backups do AWS Backup:', backup_info, 2)
add_section('Lambda Functions:', lambda_info, 2)
add_section('S3 Buckets:', s3_info, 2)

# Salva o documento e define a formatação
for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)

doc.save('informacoes_aws.docx')
print('Informações detalhadas foram adicionadas ao arquivo "informacoes_aws.docx".')

